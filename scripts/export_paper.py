
from pathlib import Path
from docx import Document
from docx.shared import Pt
import re, sys

SRC = Path('paper/paper_draft.md')
OUT = Path('paper/exports')
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / 'paper_draft.docx'
PDF = OUT / 'paper_draft.pdf'

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

lines = SRC.read_text(encoding='utf-8', errors='replace').splitlines()
in_code = False
table_lines = []

def flush_table():
    global table_lines
    if not table_lines:
        return
    rows = []
    for line in table_lines:
        if '---' in line:
            continue
        cells = [c.strip().strip('*') for c in line.strip('|').split('|')]
        rows.append(cells)
    if rows:
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = 'Table Grid'
        for i, c in enumerate(rows[0]):
            table.rows[0].cells[i].text = c
        for row in rows[1:]:
            cells = table.add_row().cells
            for i, c in enumerate(row[:len(cells)]):
                cells[i].text = c
    table_lines = []

for line in lines:
    s = line.rstrip()
    if s.startswith('```'):
        flush_table()
        in_code = not in_code
        continue
    if in_code:
        p = doc.add_paragraph()
        r = p.add_run(s)
        r.font.name = 'Consolas'
        r.font.size = Pt(8)
        continue
    if s.startswith('|') and s.endswith('|'):
        table_lines.append(s)
        continue
    flush_table()
    if not s:
        continue
    if s.startswith('# '):
        doc.add_heading(s[2:], 1)
    elif s.startswith('## '):
        doc.add_heading(s[3:], 2)
    elif s.startswith('### '):
        doc.add_heading(s[4:], 3)
    elif s.startswith('- '):
        doc.add_paragraph(s[2:], style='List Bullet')
    elif re.match(r'^\d+\.\s+', s):
        doc.add_paragraph(re.sub(r'^\d+\.\s+', '', s), style='List Number')
    else:
        doc.add_paragraph(s)

flush_table()
doc.save(DOCX)
print('Wrote', DOCX)

try:
    import win32com.client
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    d = word.Documents.Open(str(DOCX.resolve()))
    d.SaveAs(str(PDF.resolve()), FileFormat=17)
    d.Close(False)
    word.Quit()
    print('Wrote', PDF)
except Exception as e:
    print('PDF export failed. DOCX is ready. Error:', e)
